from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


import pandas as pd
import os


def customize_heading(doc, text, level, styles):
    """
    Customize a heading in the Word document based on the provided styles.

    Args:
        doc (Document): The Word document object.
        text (str): The text for the heading.
        level (int): The heading level (e.g., 1, 2, 3).
        styles (dict): A dictionary defining the styles for each level.
                       Example: {1: {"font_size": 16, "bold": True, "font_name": "Arial", "color": (0, 0, 0)}}
    """
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if level > 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = heading.runs[0]

    # Apply styles based on the level
    if level in styles:
        style = styles[level]
        run.font.size = Pt(style.get("font_size", 12))  # Default font size: 12
        run.font.name = style.get("font_name", "Arial")  # Default font: Arial
        run.bold = style.get("bold", False)  # Default bold: False
        if "color" in style:
            color = style["color"]
            run.font.color.rgb = RGBColor(*color)  # Apply color (RGB tuple)

def set_table_borders(table):
    """
    Set borders for a table to ensure gridlines are visible.
    """
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border = OxmlElement(f"w:{border_name}")
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "4")
                border.set(qn("w:space"), "0")
                border.set(qn("w:color"), "auto")
                tcBorders.append(border)
            tcPr.append(tcBorders)

def generate_docx(table_part1, table_part2,c, filename):
    doc = Document()

    # print("Final line of table_part2:")
    # print(table_part2.iloc[-1][1])  # Print the last row of table_part2

    # Check if table_part1 or table_part2 is None or empty
    if table_part1 is None or table_part2 is None or table_part1.empty or table_part2.empty:
        # Add the specified text
        paragraph_text = (
            "Hey fella, calculations have not been performed. "
            "Go back and upload a .DXF file and design the grounding grid properly before printing the outputs."
        )
        paragraph = doc.add_paragraph(paragraph_text)
        paragraph.style.font.name = "Arial"
        paragraph.style.font.size = Pt(11)

        # Save the document
        output_path = os.path.join(os.getcwd(), filename)
        doc.save(output_path)
        return output_path

    # Define styles for each heading level, including color
    heading_styles = {
        1: {"font_size": 14, "bold": True, "font_name": "Arial", "color": (0, 0, 0)},
        2: {"font_size": 12, "bold": True, "font_name": "Arial", "color": (0, 0, 0)}, 
        3: {"font_size": 12, "bold": False, "font_name": "Arial", "color": (0, 0, 0)}, 
    }

    # Check compliance from the last position of table_part2
    compliance = table_part2.iloc[-1, 1]  # Assuming compliance is in the first column of the last row
    if compliance is False or str(compliance).lower() == "false":
        # Extract relevant values for the failure message
        touch_voltage = table_part2.iloc[-3, 1]  # Assuming touch voltage is in the second column
        touch_voltage_limit = table_part2.iloc[-4, 1]  # Assuming touch voltage limit is in the third column
        step_voltage = table_part2.iloc[-2, 1]  # Assuming step voltage is in the fourth column
        step_voltage_limit = table_part2.iloc[-5, 1]  # Assuming step voltage limit is in the fifth column

        # Write the failure message (Part 1)
        failure_text_part1 = (
            "You lose fella! The grounding grid is not complying with a safe design. Go back to the design."
        )
        failure_paragraph_part1 = doc.add_paragraph(failure_text_part1)
        failure_paragraph_part1.style.font.name = "Arial"
        failure_paragraph_part1.style.font.size = Pt(11)

        # Write the failure message (Part 2)
        failure_text_part2 = (
            "The Touch voltage is "
        )
        failure_paragraph_part2 = doc.add_paragraph(failure_text_part2)
        failure_paragraph_part2.style.font.name = "Arial"
        failure_paragraph_part2.style.font.size = Pt(11)

        # Add touch voltage in bold and red
        touch_run = failure_paragraph_part2.add_run(f"{touch_voltage} ")
        touch_run.bold = True
        touch_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color

        # Continue the text
        failure_paragraph_part2.add_run("and the tolerable touch voltage is ").font.size = Pt(11)
        touch_limit_run = failure_paragraph_part2.add_run(f"{touch_voltage_limit} ")
        touch_limit_run.bold = True
        touch_limit_run.font.color.rgb = RGBColor(0, 0, 0)  # Black color
        failure_paragraph_part2.add_run(
            f"On the other hand, the step voltage is {step_voltage} "
            f"and the tolerable step voltage is {step_voltage_limit}."
        ).font.size = Pt(11)

        # Save the document
        output_path = os.path.join(os.getcwd(), filename)
        doc.save(output_path)
        return output_path

    # Add and customize the main heading (Level 1)
    customize_heading(doc, "Grounding Grid Calculation Results", level=1, styles=heading_styles)

    # Add a new section for "Design Considerations and Assumptions"
    customize_heading(doc, "Design Considerations and Assumptions", level=1, styles=heading_styles)

    # Add the paragraph with dynamic values from the `c` dictionary
    resistivity = c.get("Soil Resistivity", "N/A")
    short_circuit_conductor = c.get("Conductor Short Circuit", "N/A")
    short_circuit = c.get("Short Circuit Current", "N/A")
    split_factor = c.get("Split Factor", "N/A")
    depth = c.get("Conductors Depth", "N/A")
    print("Depth:", depth)
    crushed_rock_resistivity = c.get("Crushed Rock Resistivity", "N/A")
    crushed_rock_depth = c.get("Crushed Rock Depth", "N/A")

    if split_factor == 1:
        split_factor = "1 (Conservative approach)"

    if short_circuit_conductor == short_circuit_conductor:
        paragraph_text_1 = (
            f"The Grounding grid for this bid has been designed following the requirements of IEEE 80, "
            f"considering a resistivity value of {resistivity}, and a maximum short-circuit current of {short_circuit_conductor} through the conductors "
            f"and towards the touch and step potential calculation. A split factor of {split_factor} was also contemplated."
        )

        paragraph_text_2 = (
            f"The grounding grid considered is shown in the drawing attached with the bid, with a depth of {depth}, "
            f"and a crushed rock resistivity of {crushed_rock_resistivity}, extending to a depth of {crushed_rock_depth}."
        )
    else:
        paragraph_text_1 = (
            f"The Grounding grid for this bid has been designed following the requirements of IEEE 80, "
            f"considering a resistivity value of {resistivity}, a maximum short-circuit current through the conductors "
            f"of {short_circuit_conductor} and a maximum short-circuit current towards the touch and step potentials calculation "
            f"of {short_circuit}. A split factor of {split_factor} was also contemplated."
        )

        paragraph_text_2 = (
            f"The grounding grid considered is shown in the drawing attached with the bid, with a depth of {depth}, "
            f"and a crushed rock resistivity of {crushed_rock_resistivity}, extending to a depth of {crushed_rock_depth}."
        )

    # Add the paragraphs to the document
    paragraph_1 = doc.add_paragraph(paragraph_text_1)
    paragraph_1.style.font.name = "Arial"
    paragraph_1.style.font.size = Pt(11)

    paragraph_2 = doc.add_paragraph(paragraph_text_2)
    paragraph_2.style.font.name = "Arial"
    paragraph_2.style.font.size = Pt(11)

    # Add an empty paragraph after paragraph 2
    doc.add_paragraph()


    def add_table_from_dict(data: dict, title: str):
            # Add and customize a subheading (Level 2)
            customize_heading(doc, title, level=2, styles=heading_styles)

            # Add a table from the dictionary
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Parameter"
            hdr_cells[1].text = "Value"
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[1].paragraphs[0].runs[0].font.bold = True
            hdr_cells[0].paragraphs[0].runs[0].font.name = "Arial"
            hdr_cells[1].paragraphs[0].runs[0].font.name = "Arial"

            # Exclude specific keys
            excluded_keys = {"Filepath", "DXF Drawing Units", "Grounding Resistance Model"}
            for key, value in data.items():
                if key not in excluded_keys:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(key)
                    row_cells[1].text = str(value)
                    row_cells[0].paragraphs[0].runs[0].font.name = "Arial"
                    row_cells[1].paragraphs[0].runs[0].font.name = "Arial"



    # Merge the two tables into one DataFrame
    merged_table = pd.concat([table_part1, table_part2], ignore_index=True)

    def add_table_from_df(df: pd.DataFrame, title: str):
        # Add and customize a subheading (Level 2)
        customize_heading(doc, title, level=2, styles=heading_styles)

        # Exclude the last row from the DataFrame
        df = df.iloc[:-1]

        # Add a table from the DataFrame
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0) 
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True  # Bold headers
            hdr_cells[i].paragraphs[0].runs[0].font.name = "Arial"  # Set font to Arial
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(11)  # Set font size to 11pt

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)
                row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 0)
                row_cells[i].paragraphs[0].runs[0].font.name = "Arial"  # Set font to Arial 
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)  # Set font size to 11pt
        
        # Format table
        set_table_borders(table)

    
    # Add the inputs table
    add_table_from_dict(c, "Inputs for Calculation")

    # Add the merged table with a heading
    add_table_from_df(merged_table, "Calculation Results")

    # Save the document
    output_path = os.path.join(os.getcwd(), filename)
    doc.save(output_path)
    return output_path